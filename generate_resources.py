
import os
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt

# Output directory
OUTPUT_DIR = "/Users/yusuf/Downloads/birstbd_foxmenstudio/bristbd_foxmenstudio_frontend/public/resources"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def create_pdf(filename, title, content):
    c = canvas.Canvas(os.path.join(OUTPUT_DIR, filename), pagesize=letter)
    width, height = letter
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, height - 72, title)
    
    c.setFont("Helvetica", 12)
    y_position = height - 100
    for line in content.split('\n'):
        if y_position < 72:
            c.showPage()
            y_position = height - 72
        c.drawString(72, y_position, line)
        y_position -= 14
        
    c.save()
    print(f"Created {filename}")

def create_docx(filename, title, content):
    doc = Document()
    doc.add_heading(title, 0)
    for line in content.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    doc.save(os.path.join(OUTPUT_DIR, filename))
    print(f"Created {filename}")

def create_excel(filename, data):
    df = pd.DataFrame(data)
    df.to_excel(os.path.join(OUTPUT_DIR, filename), index=False)
    print(f"Created {filename}")

def create_csv(filename, data):
    df = pd.DataFrame(data)
    df.to_csv(os.path.join(OUTPUT_DIR, filename), index=False)
    print(f"Created {filename}")

# 1. Thesis Template (APA Style) - DOCX
thesis_content = """
Abstract
[Your Abstract Here]

Introduction
[Introduction to the topic...]

Literature Review
[Review of related literature...]

Methodology
[Research design, participants, measures...]

Results
[Findings...]

Discussion
[Interpretation of results...]

References
[APA Style References...]
"""
create_docx("Thesis_Template_APA.docx", "Thesis Template (APA Style)", thesis_content)

# 2. Research Proposal Sample - PDF
proposal_content = """
Title: Impact of AI on Academic integrity

1. Introduction
The rapid advancement of Artificial Intelligence (AI) has transformed...

2. Problem Statement
Current academic institutions lack standardized frameworks...

3. Research Questions
- to what extent does AI assistance influence student writing?
- What correspond with the detectable patterns?

4. Methodology
Mixed-methods approach utilizing surveys...

5. Significance
This study will contribute to...
"""
create_pdf("Research_Proposal_Sample.pdf", "Research Proposal Sample", proposal_content)

# 3. Statistical Test Cheat Sheet - PDF
cheat_sheet_content = """
1. Comparing Means
- 2 Groups (Independent): Independent t-test
- 2 Groups (Paired): Paired t-test
- 3+ Groups: ANOVA

2. Relationship
- 2 Continuous Variables: Pearson Correlation
- 2 Ordinal/Non-normal: Spearman Correlation

3. Prediction
- Continuous Outcome: Linear Regression
- Binary Outcome: Logistic Regression
"""
create_pdf("Statistical_Test_Cheat_Sheet.pdf", "Statistical Test Cheat Sheet", cheat_sheet_content)

# 4. Sample Dataset (SPSS/CSV) - CSV
data = {
    "ID": [1, 2, 3, 4, 5],
    "Age": [23, 25, 22, 24, 26],
    "Gender": ["M", "F", "F", "M", "M"],
    "Score_Pre": [70, 85, 78, 65, 80],
    "Score_Post": [75, 88, 82, 70, 85]
}
create_csv("Sample_Dataset.csv", data)

# 5. Literature Review Matrix - XLSX
matrix_data = {
    "Author": ["Smith (2020)", "Doe (2021)"],
    "Title": ["Study A", "Study B"],
    "Methodology": ["Survey", "Experiment"],
    "Key Findings": ["Result X", "Result Y"],
    "Gaps": ["Sample size", "Context"]
}
create_excel("Literature_Review_Matrix.xlsx", matrix_data)

# 6. Consent Form Template - DOCX
consent_content = """
Participant Consent Form

Title of Study: [Study Title]
Principal Investigator: [Name]

I have read the information sheet and understand the purpose of this study.
I understand that my participation is voluntary.
I agree to take part in this study.

Name: __________________________
Signature: _______________________
Date: ___________________________
"""
create_docx("Consent_Form_Template.docx", "Consent Form Template", consent_content)
